from __future__ import annotations

import csv
import json
from pathlib import Path
from typing import Any

from openpyxl import load_workbook

TEXT_EXTENSIONS = {".md", ".txt", ".log"}
CSV_MAX_ROWS = 30
EXCEL_MAX_SHEETS = 25
EXCEL_MAX_ROWS = 30
EXCEL_MAX_COLUMNS = 20
PDF_MAX_PAGES = 20
DOCX_MAX_TABLES = 10
TEXT_CAP = 12_000


def extract_document(path: Path | str) -> dict[str, Any]:
    source = Path(path)
    extension = source.suffix.lower()

    if extension in {".doc", ".xls"}:
        return _result(source, "limited", "", reason="legacy binary format skipped")
    if extension in TEXT_EXTENSIONS:
        return _result(source, "ok", _cap_text(_read_text_file(source)))
    if extension == ".csv":
        return _result(source, "ok", _cap_text(_extract_csv(source)))
    if extension == ".json":
        return _result(source, "ok", _cap_text(_extract_json(source)))
    if extension in {".xlsx", ".xlsm"}:
        return _result(source, "ok", _cap_text(_extract_excel(source)))
    if extension == ".pdf":
        return _extract_pdf(source)
    if extension == ".docx":
        return _extract_docx(source)

    return _result(source, "limited", "", reason="unsupported extension")


def _result(path: Path, status: str, text: str, reason: str = "") -> dict[str, Any]:
    payload: dict[str, Any] = {
        "status": status,
        "text": text,
        "path": str(path),
        "extension": path.suffix.lower(),
    }
    if reason:
        payload["reason"] = reason
    return payload


def _read_text_file(path: Path) -> str:
    data = path.read_bytes()
    for encoding in ("utf-8", "utf-8-sig", "cp949"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _extract_csv(path: Path) -> str:
    text = _read_text_file(path)
    rows: list[str] = []
    reader = csv.reader(text.splitlines())
    for index, row in enumerate(reader):
        if index >= CSV_MAX_ROWS:
            break
        rows.append(", ".join(row))
    return "\n".join(rows)


def _extract_json(path: Path) -> str:
    text = _read_text_file(path)
    try:
        parsed = json.loads(text)
    except json.JSONDecodeError:
        return text
    return json.dumps(parsed, ensure_ascii=False, indent=2)


def _extract_excel(path: Path) -> str:
    workbook = load_workbook(path, read_only=True, data_only=True)
    try:
        parts: list[str] = []
        for sheet_index, worksheet in enumerate(workbook.worksheets):
            if sheet_index >= EXCEL_MAX_SHEETS:
                break
            parts.append(f"# Sheet: {worksheet.title}")
            for row_index, row in enumerate(
                worksheet.iter_rows(
                    min_row=1,
                    max_row=EXCEL_MAX_ROWS,
                    max_col=EXCEL_MAX_COLUMNS,
                    values_only=True,
                )
            ):
                if row_index >= EXCEL_MAX_ROWS:
                    break
                values = ["" if value is None else str(value) for value in row]
                if any(values):
                    parts.append("\t".join(values).rstrip())
        return "\n".join(parts)
    finally:
        workbook.close()


def _extract_pdf(path: Path) -> dict[str, Any]:
    try:
        from pypdf import PdfReader
    except ImportError:
        return _result(path, "limited", "", reason="pypdf is not installed")

    reader = PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages[:PDF_MAX_PAGES]:
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(page_text)
    text = _cap_text("\n\n".join(pages))
    if not text.strip():
        return _result(path, "limited", "", reason="no extractable pdf text")
    return _result(path, "ok", text)


def _extract_docx(path: Path) -> dict[str, Any]:
    try:
        from docx import Document
    except ImportError:
        return _result(path, "limited", "", reason="python-docx is not installed")

    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables[:DOCX_MAX_TABLES]:
        for row in table.rows:
            parts.append("\t".join(cell.text.strip() for cell in row.cells))
    return _result(path, "ok", _cap_text("\n".join(parts)))


def _cap_text(text: str) -> str:
    if len(text) <= TEXT_CAP:
        return text
    return text[:TEXT_CAP]
